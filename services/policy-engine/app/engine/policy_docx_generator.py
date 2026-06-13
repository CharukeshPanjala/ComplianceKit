# WHAT: Policy DOCX generator | CHANGE: new file | WHY: COM-176 — render a policy's Markdown content as a Word document for download

import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from common.models.policy import Policy

from app.engine.markdown_blocks import (
    parse_markdown,
    Span,
    Heading as MdHeading,
    Paragraph as MdParagraph,
    ListItem as MdListItem,
    Table as MdTable,
)


# ── Colours ───────────────────────────────────────────────────────────────────

NAVY = RGBColor(0x0F, 0x20, 0x44)
TEXT_GRAY = RGBColor(0x37, 0x41, 0x51)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_spans(paragraph, spans: list[Span]) -> None:
    for sp in spans:
        run = paragraph.add_run(sp.text)
        run.bold = sp.bold
        run.italic = sp.italic


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _render_table(doc: Document, block: MdTable) -> None:
    num_cols = len(block.headers) or 1
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, cell_spans in enumerate(block.headers):
        p = header_cells[idx].paragraphs[0]
        for sp in cell_spans:
            run = p.add_run(sp.text)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(header_cells[idx], "0F2044")

    for row in block.rows:
        cells = table.add_row().cells
        for idx, cell_spans in enumerate(row):
            if idx >= num_cols:
                break
            _add_spans(cells[idx].paragraphs[0], cell_spans)


# ── Generic markdown renderer ─────────────────────────────────────────────────

def render_markdown_docx(title: str, subtitle: str, content: str) -> bytes:
    """Render a title, subtitle and Markdown body as a Word document."""
    doc = Document()

    title_para = doc.add_heading("", level=0)
    title_run = title_para.add_run(title)
    title_run.font.color.rgb = NAVY

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = TEXT_GRAY

    blocks = parse_markdown(content or "")
    skipped_h1 = False

    for block in blocks:
        if isinstance(block, MdHeading):
            if block.level == 1 and not skipped_h1:
                skipped_h1 = True
                continue
            heading = doc.add_heading("", level=min(block.level, 3))
            _add_spans(heading, block.spans)
        elif isinstance(block, MdParagraph):
            _add_spans(doc.add_paragraph(), block.spans)
        elif isinstance(block, MdListItem):
            style = "List Number" if block.ordered else "List Bullet"
            _add_spans(doc.add_paragraph(style=style), block.spans)
        elif isinstance(block, MdTable):
            _render_table(doc, block)
            doc.add_paragraph()

    if not blocks:
        doc.add_paragraph("This document has no content yet.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Main export function ──────────────────────────────────────────────────────

def generate_policy_docx(policy: Policy) -> bytes:
    status_label = policy.status.value.replace("_", " ").title() if policy.status else "Draft"
    title = policy.title or "Policy"
    if policy.tenant_name:
        title = f"{title} — {policy.tenant_name}"
    subtitle = (
        f"Generated {date.today().strftime('%d %B %Y')}  ·  "
        f"Version {policy.current_version}  ·  {status_label}"
    )
    return render_markdown_docx(title, subtitle, policy.content or "")
