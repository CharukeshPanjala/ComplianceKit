# WHAT: Tests for compliance report generator + API | CHANGE: new file | WHY: COM-177 — cover executive summary generation, markdown assembly, and the /api/v1/reports/compliance endpoint
"""
Tests for:
- app.engine.compliance_report_generator (ComplianceReportGenerator)
- GET /api/v1/reports/compliance
"""
import io
import uuid
from datetime import datetime, timezone
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.main import app
from app.config import PolicySettings
from app.engine.compliance_report_generator import (
    ComplianceReportGenerator,
    GapSummary,
    RegulationSummary,
    ProfileSummary,
)
from common.auth.clerk import verify_token, TokenClaims
from common.db.session import get_admin_session
from common.models.assessment import (
    Assessment, Gap, AssessmentStatus, RiskLevel, GapStatus, RemediationPriority,
)
from common.models.company_profile import CompanyProfile, Industry, CompanySize, B2BOrB2C, DataRole
from common.models.regulation import Regulation


# ── Settings ──────────────────────────────────────────────────────────────────

AI_DISABLED_SETTINGS = PolicySettings(azure_openai_endpoint="", azure_openai_api_key="")
AI_ENABLED_SETTINGS = PolicySettings(
    azure_openai_endpoint="https://test.openai.azure.com", azure_openai_api_key="test-key"
)


# ── Helpers (dataclass fixtures) ─────────────────────────────────────────────

def make_profile_summary(**kwargs) -> ProfileSummary:
    return ProfileSummary(
        tenant_name=kwargs.get("tenant_name", "Acme Inc"),
        industry=kwargs.get("industry", "technology"),
        company_size=kwargs.get("company_size", "11-50"),
        b2b_or_b2c=kwargs.get("b2b_or_b2c", "b2b"),
        data_role=kwargs.get("data_role", "controller"),
        primary_jurisdiction=kwargs.get("primary_jurisdiction", "IE"),
        has_compliance_officer=kwargs.get("has_compliance_officer", True),
        dpo_name=kwargs.get("dpo_name", "Jane Smith"),
    )


def make_gap_summary(**kwargs) -> GapSummary:
    return GapSummary(
        article=kwargs.get("article", "Article 32"),
        title=kwargs.get("title", "No encryption policy"),
        severity=kwargs.get("severity", "critical"),
        status=kwargs.get("status", "not_met"),
        remediation_priority=kwargs.get("remediation_priority", "critical"),
        remediation_hint=kwargs.get("remediation_hint", "Implement encryption at rest"),
        resolved=kwargs.get("resolved", False),
    )


def make_regulation_summary(**kwargs) -> RegulationSummary:
    return RegulationSummary(
        regulation_name=kwargs.get("regulation_name", "GDPR"),
        status=kwargs.get("status", "completed"),
        score=kwargs.get("score", 72),
        risk_level=kwargs.get("risk_level", "medium"),
        applicable_rules=kwargs.get("applicable_rules", 10),
        met_rules=kwargs.get("met_rules", 5),
        partial_rules=kwargs.get("partial_rules", 2),
        not_met_rules=kwargs.get("not_met_rules", 2),
        unknown_rules=kwargs.get("unknown_rules", 1),
        completed_at=kwargs.get("completed_at", datetime.now(timezone.utc)),
        gaps=kwargs.get("gaps", []),
    )


def make_never_run_summary(regulation_name: str) -> RegulationSummary:
    return RegulationSummary(
        regulation_name=regulation_name,
        status="never_run",
        score=None,
        risk_level=None,
        applicable_rules=0,
        met_rules=0,
        partial_rules=0,
        not_met_rules=0,
        unknown_rules=0,
        completed_at=None,
        gaps=[],
    )


# ═══════════════════════════════════════════════════════════════════════════════
# ComplianceReportGenerator — executive summary
# ═══════════════════════════════════════════════════════════════════════════════

class TestExecutiveSummary:

    def test_stub_summary_has_three_paragraphs(self):
        gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)
        text = gen._stub_summary(make_profile_summary(), [make_regulation_summary()], 72)

        paragraphs = text.split("\n\n")
        assert len(paragraphs) == 3
        assert "72%" in paragraphs[0]

    def test_stub_summary_notes_never_run_regulations(self):
        gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)
        regs = [make_regulation_summary(regulation_name="GDPR"), make_never_run_summary("NIS2")]
        text = gen._stub_summary(make_profile_summary(), regs, 72)

        assert "NIS2" in text
        assert "have not yet been run" in text

    def test_stub_summary_highlights_critical_gaps(self):
        gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)
        critical_gap = make_gap_summary(article="Article 32", title="No encryption policy")
        regs = [make_regulation_summary(gaps=[critical_gap])]
        text = gen._stub_summary(make_profile_summary(), regs, 72)

        assert "Article 32" in text
        assert "No encryption policy" in text

    def test_stub_summary_no_critical_gaps(self):
        gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)
        met_gap = make_gap_summary(status="met", remediation_priority=None)
        regs = [make_regulation_summary(gaps=[met_gap])]
        text = gen._stub_summary(make_profile_summary(), regs, 100)

        assert "no outstanding critical-priority gaps" in text

    @pytest.mark.asyncio
    async def test_ai_disabled_returns_stub(self):
        gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)
        text, is_ai = await gen.generate_executive_summary(make_profile_summary(), [make_regulation_summary()], 72)

        assert is_ai is False
        assert "72%" in text

    @pytest.mark.asyncio
    @patch("app.engine.compliance_report_generator.get_async_client")
    async def test_ai_enabled_returns_ai_text(self, mock_get_client):
        mock_response = MagicMock()
        mock_response.choices = [MagicMock(message=MagicMock(content="AI-written executive summary."))]
        mock_client = MagicMock()
        mock_client.chat.completions.create = AsyncMock(return_value=mock_response)
        mock_get_client.return_value = mock_client

        gen = ComplianceReportGenerator(AI_ENABLED_SETTINGS)
        text, is_ai = await gen.generate_executive_summary(make_profile_summary(), [make_regulation_summary()], 72)

        assert is_ai is True
        assert text == "AI-written executive summary."

    @pytest.mark.asyncio
    @patch("app.engine.compliance_report_generator.get_async_client")
    async def test_ai_failure_falls_back_to_stub(self, mock_get_client):
        mock_get_client.side_effect = Exception("Azure unavailable")

        gen = ComplianceReportGenerator(AI_ENABLED_SETTINGS)
        text, is_ai = await gen.generate_executive_summary(make_profile_summary(), [make_regulation_summary()], 72)

        assert is_ai is False
        assert "72%" in text


# ═══════════════════════════════════════════════════════════════════════════════
# ComplianceReportGenerator — build_report_markdown
# ═══════════════════════════════════════════════════════════════════════════════

class TestBuildReportMarkdown:

    def setup_method(self):
        self.gen = ComplianceReportGenerator(AI_DISABLED_SETTINGS)

    def test_includes_all_top_level_sections(self):
        regs = [
            make_regulation_summary(regulation_name="GDPR", gaps=[make_gap_summary()]),
            make_regulation_summary(regulation_name="NIS2", score=60, risk_level="high", gaps=[]),
            make_never_run_summary("EU_AI_ACT"),
        ]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 66, "Exec summary text.")

        assert "# Compliance Report" in md
        assert "## Executive Summary" in md
        assert "Exec summary text." in md
        assert "## Regulation Breakdown" in md
        assert "## Gap Analysis — GDPR" in md
        assert "## Gap Analysis — NIS2" in md
        assert "## Gap Analysis — EU AI Act" in md
        assert "## Remediation Roadmap" in md
        assert "## Profile Summary" in md

    def test_regulation_breakdown_table_rows(self):
        regs = [make_regulation_summary(regulation_name="GDPR", score=72, risk_level="medium")]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 72, "Summary.")

        assert "| GDPR | Completed | 72% | Medium | 10 | 5 | 2 | 2 | 1 |" in md

    def test_never_run_regulation_breakdown_row(self):
        regs = [make_never_run_summary("NIS2")]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 72, "Summary.")

        assert "| NIS2 | Never Run | — | — | 0 | 0 | 0 | 0 | 0 |" in md
        assert "_No completed assessment for this regulation yet._" in md

    def test_gap_analysis_table_includes_met_gaps(self):
        met_gap = make_gap_summary(article="Article 5", title="Lawful basis documented", status="met",
                                     remediation_priority=None)
        not_met_gap = make_gap_summary(article="Article 32", title="No encryption policy")
        regs = [make_regulation_summary(gaps=[met_gap, not_met_gap])]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 72, "Summary.")

        assert "| Article 5 | Lawful basis documented | Critical | Met | — |" in md
        assert "| Article 32 | No encryption policy | Critical | Not Met | Critical |" in md

    def test_no_applicable_rules_message(self):
        regs = [make_regulation_summary(gaps=[])]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 72, "Summary.")

        assert "_No applicable rules found for this regulation._" in md

    def test_remediation_roadmap_groups_by_priority(self):
        critical_gap = make_gap_summary(article="Article 32", remediation_priority="critical")
        low_gap = make_gap_summary(article="Article 13", remediation_priority="low", severity="low")
        regs = [make_regulation_summary(gaps=[critical_gap, low_gap])]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 72, "Summary.")

        assert "### Critical Priority" in md
        assert "### Low Priority" in md
        critical_idx = md.index("### Critical Priority")
        low_idx = md.index("### Low Priority")
        assert critical_idx < low_idx
        assert "[GDPR Article 32]" in md

    def test_remediation_roadmap_excludes_resolved_and_met(self):
        resolved_gap = make_gap_summary(article="Article 30", resolved=True)
        met_gap = make_gap_summary(article="Article 5", status="met", remediation_priority=None)
        regs = [make_regulation_summary(gaps=[resolved_gap, met_gap])]
        md = self.gen.build_report_markdown(make_profile_summary(), regs, 100, "Summary.")

        assert "_No outstanding remediation items — all applicable gaps are met or resolved._" in md

    def test_profile_summary_fields(self):
        profile = make_profile_summary(
            tenant_name="Acme Inc", industry="technology", company_size="11-50",
            b2b_or_b2c="b2b", data_role="controller", primary_jurisdiction="IE",
            has_compliance_officer=True, dpo_name="Jane Smith",
        )
        md = self.gen.build_report_markdown(profile, [make_regulation_summary()], 72, "Summary.")

        assert "- **Company:** Acme Inc" in md
        assert "- **Industry:** Technology" in md
        assert "- **Company Size:** 11-50" in md
        assert "- **Business Model:** B2B" in md
        assert "- **Data Role:** Controller" in md
        assert "- **Primary Jurisdiction:** IE" in md
        assert "- **Compliance Officer Appointed:** Yes" in md
        assert "- **DPO:** Jane Smith" in md


# ═══════════════════════════════════════════════════════════════════════════════
# GET /api/v1/reports/compliance
# ═══════════════════════════════════════════════════════════════════════════════

FAKE_CLAIMS = TokenClaims(user_id="user_test123", tenant_id="ten_test456", org_role="org:admin")


def scalar_result(value):
    r = MagicMock()
    r.scalar_one_or_none.return_value = value
    return r


def scalars_result(values):
    scalars = MagicMock()
    scalars.all.return_value = values
    r = MagicMock()
    r.scalars.return_value = scalars
    return r


def make_mock_profile(**kwargs):
    p = MagicMock(spec=CompanyProfile)
    p.tenant_id = FAKE_CLAIMS.tenant_id
    p.tenant_name = kwargs.get("tenant_name", "Acme Inc")
    p.industry = kwargs.get("industry", Industry.TECHNOLOGY)
    p.company_size = kwargs.get("company_size", CompanySize.ELEVEN_TO_FIFTY)
    p.b2b_or_b2c = kwargs.get("b2b_or_b2c", B2BOrB2C.B2B)
    p.data_role = kwargs.get("data_role", DataRole.CONTROLLER)
    p.primary_jurisdiction = kwargs.get("primary_jurisdiction", "IE")
    p.has_compliance_officer = kwargs.get("has_compliance_officer", True)
    p.dpo_name = kwargs.get("dpo_name", "Jane Smith")
    return p


def make_mock_regulation(name: str):
    r = MagicMock(spec=Regulation)
    r.id = uuid.uuid4()
    r.name = name
    return r


def make_mock_assessment(**kwargs):
    a = MagicMock(spec=Assessment)
    a.id = uuid.uuid4()
    a.status = kwargs.get("status", AssessmentStatus.COMPLETED)
    a.score = kwargs.get("score", 72)
    a.risk_level = kwargs.get("risk_level", RiskLevel.MEDIUM)
    a.applicable_rules = kwargs.get("applicable_rules", 10)
    a.met_rules = kwargs.get("met_rules", 5)
    a.partial_rules = kwargs.get("partial_rules", 2)
    a.not_met_rules = kwargs.get("not_met_rules", 2)
    a.unknown_rules = kwargs.get("unknown_rules", 1)
    a.completed_at = kwargs.get("completed_at", datetime.now(timezone.utc))
    return a


def make_mock_gap(**kwargs):
    g = MagicMock(spec=Gap)
    g.article = kwargs.get("article", "Article 32")
    g.article_number = kwargs.get("article_number", 32)
    g.title = kwargs.get("title", "No encryption policy")
    g.severity = kwargs.get("severity", "critical")
    g.status = kwargs.get("status", GapStatus.NOT_MET)
    g.remediation_priority = kwargs.get("remediation_priority", RemediationPriority.CRITICAL)
    g.remediation_hint = kwargs.get("remediation_hint", "Implement encryption at rest")
    g.resolved = kwargs.get("resolved", False)
    return g


def build_side_effect(profile, reg_configs):
    """reg_configs: list of (regulation_or_none, assessment_or_none, gaps_or_none)."""
    responses = [scalar_result(profile)]
    for regulation, assessment, gaps in reg_configs:
        responses.append(scalar_result(regulation))
        if regulation is None:
            continue
        responses.append(scalar_result(assessment))
        if assessment is not None and assessment.status == AssessmentStatus.COMPLETED:
            responses.append(scalars_result(gaps or []))

    responses_iter = iter(responses)

    async def side_effect(stmt):
        return next(responses_iter)

    return side_effect


def real_generator_for_markdown(*args, **kwargs):
    return ComplianceReportGenerator(AI_DISABLED_SETTINGS).build_report_markdown(*args, **kwargs)


@pytest.fixture
def client():
    mock_session = AsyncMock()
    mock_session.__aenter__ = AsyncMock(return_value=mock_session)
    mock_session.__aexit__ = AsyncMock(return_value=None)

    async def override_session():
        yield mock_session

    async def override_token():
        return FAKE_CLAIMS

    app.dependency_overrides[get_admin_session] = override_session
    app.dependency_overrides[verify_token] = override_token
    yield TestClient(app), mock_session
    app.dependency_overrides.clear()


@pytest.fixture
def mock_generator():
    """Patches ComplianceReportGenerator in the router: stubs the AI summary (no Azure
    calls) but delegates build_report_markdown to a real instance so PDF/DOCX
    rendering is exercised end-to-end."""
    with patch("app.routers.reports.ComplianceReportGenerator") as mock_cls:
        instance = mock_cls.return_value
        instance.generate_executive_summary = AsyncMock(
            return_value=("Stub executive summary for tests.", False)
        )
        instance.build_report_markdown.side_effect = real_generator_for_markdown
        yield mock_cls


class TestComplianceReportEndpoint:

    def test_returns_pdf_for_all_completed(self, client, mock_generator):
        test_client, mock_session = client
        profile = make_mock_profile()
        reg_configs = [
            (make_mock_regulation("GDPR"), make_mock_assessment(score=72, risk_level=RiskLevel.MEDIUM),
             [make_mock_gap(article="Article 32", status=GapStatus.NOT_MET)]),
            (make_mock_regulation("NIS2"), make_mock_assessment(score=60, risk_level=RiskLevel.HIGH), []),
            (make_mock_regulation("EU_AI_ACT"), make_mock_assessment(score=85, risk_level=RiskLevel.LOW), []),
        ]
        mock_session.execute = build_side_effect(profile, reg_configs)

        response = test_client.get("/api/v1/reports/compliance")

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/pdf"
        assert "attachment; filename=compliance_report.pdf" in response.headers["content-disposition"]
        assert response.content.startswith(b"%PDF")

    def test_returns_docx_with_expected_sections(self, client, mock_generator):
        test_client, mock_session = client
        profile = make_mock_profile()
        reg_configs = [
            (make_mock_regulation("GDPR"), make_mock_assessment(score=72, risk_level=RiskLevel.MEDIUM),
             [make_mock_gap(article="Article 32", status=GapStatus.NOT_MET)]),
            (make_mock_regulation("NIS2"), make_mock_assessment(score=60, risk_level=RiskLevel.HIGH), []),
            (make_mock_regulation("EU_AI_ACT"), make_mock_assessment(score=85, risk_level=RiskLevel.LOW), []),
        ]
        mock_session.execute = build_side_effect(profile, reg_configs)

        response = test_client.get("/api/v1/reports/compliance?format=docx")

        assert response.status_code == 200
        assert response.headers["content-type"] == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert "attachment; filename=compliance_report.docx" in response.headers["content-disposition"]
        assert response.content.startswith(b"PK")

        doc = DocxDocument(io.BytesIO(response.content))
        assert "Acme Inc" in doc.paragraphs[0].text
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Executive Summary" in headings
        assert "Regulation Breakdown" in headings
        assert "Gap Analysis — GDPR" in headings
        assert "Remediation Roadmap" in headings
        assert "Profile Summary" in headings

    def test_partial_report_overall_score_is_average_of_completed(self, client, mock_generator):
        test_client, mock_session = client
        profile = make_mock_profile()
        reg_configs = [
            (make_mock_regulation("GDPR"), make_mock_assessment(score=80, risk_level=RiskLevel.LOW), []),
            (make_mock_regulation("NIS2"), make_mock_assessment(score=60, risk_level=RiskLevel.HIGH), []),
            (make_mock_regulation("EU_AI_ACT"), None, None),
        ]
        mock_session.execute = build_side_effect(profile, reg_configs)

        response = test_client.get("/api/v1/reports/compliance?format=docx")

        assert response.status_code == 200
        doc = DocxDocument(io.BytesIO(response.content))
        assert "Overall Score 70%" in doc.paragraphs[1].text

        breakdown_table = doc.tables[0]
        row_texts = ["|".join(cell.text for cell in row.cells) for row in breakdown_table.rows]
        assert any("NIS2" in t and "60%" in t and "High" in t for t in row_texts)
        assert any("EU AI Act" in t and "Never Run" in t for t in row_texts)

    def test_422_when_no_completed_assessments(self, client, mock_generator):
        test_client, mock_session = client
        profile = make_mock_profile()
        reg_configs = [
            (make_mock_regulation("GDPR"), None, None),
            (make_mock_regulation("NIS2"), None, None),
            (make_mock_regulation("EU_AI_ACT"), None, None),
        ]
        mock_session.execute = build_side_effect(profile, reg_configs)

        response = test_client.get("/api/v1/reports/compliance")

        assert response.status_code == 422
        assert "No completed assessments yet" in response.json()["detail"]

    def test_404_when_profile_missing(self, client, mock_generator):
        test_client, mock_session = client
        mock_session.execute = AsyncMock(return_value=scalar_result(None))

        response = test_client.get("/api/v1/reports/compliance")

        assert response.status_code == 404

    def test_invalid_format_returns_422(self, client, mock_generator):
        test_client, _ = client

        response = test_client.get("/api/v1/reports/compliance?format=xyz")

        assert response.status_code == 422
