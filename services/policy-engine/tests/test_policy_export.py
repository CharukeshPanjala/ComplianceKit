# WHAT: Tests for policy export engine | CHANGE: new file | WHY: COM-176 — cover Markdown block parsing and PDF/DOCX rendering for policy documents
"""
Unit tests for:
- app.engine.markdown_blocks (parse_inline, parse_markdown)
- app.engine.policy_pdf_generator (generate_policy_pdf)
- app.engine.policy_docx_generator (generate_policy_docx)
"""
import io

from docx import Document as DocxDocument

from common.models.policy import Policy, PolicyType, PolicyStatus, PolicyContentFormat

from app.engine.markdown_blocks import (
    parse_inline, parse_markdown, Span, Heading, Paragraph, ListItem, Table,
)
from app.engine.policy_pdf_generator import generate_policy_pdf
from app.engine.policy_docx_generator import generate_policy_docx


# ═══════════════════════════════════════════════════════════════════════════════
# parse_inline
# ═══════════════════════════════════════════════════════════════════════════════

class TestParseInline:

    def test_plain_text(self):
        assert parse_inline("Hello world") == [Span("Hello world")]

    def test_bold(self):
        assert parse_inline("**bold**") == [Span("bold", bold=True)]

    def test_italic(self):
        assert parse_inline("_italic_") == [Span("italic", italic=True)]

    def test_mixed(self):
        assert parse_inline("normal **bold** and _italic_ text") == [
            Span("normal "),
            Span("bold", bold=True),
            Span(" and "),
            Span("italic", italic=True),
            Span(" text"),
        ]


# ═══════════════════════════════════════════════════════════════════════════════
# parse_markdown
# ═══════════════════════════════════════════════════════════════════════════════

class TestParseMarkdown:

    def test_headings(self):
        blocks = parse_markdown("# Title\n\n## Section\n\n### Sub")
        assert blocks == [
            Heading(level=1, spans=[Span("Title")]),
            Heading(level=2, spans=[Span("Section")]),
            Heading(level=3, spans=[Span("Sub")]),
        ]

    def test_paragraph_joins_wrapped_lines(self):
        blocks = parse_markdown("This is a paragraph\nspanning two lines.")
        assert blocks == [Paragraph(spans=parse_inline("This is a paragraph spanning two lines."))]

    def test_bullet_list(self):
        blocks = parse_markdown("- one\n- two")
        assert blocks == [
            ListItem(spans=[Span("one")], ordered=False),
            ListItem(spans=[Span("two")], ordered=False),
        ]

    def test_ordered_list(self):
        blocks = parse_markdown("1. first\n2. second")
        assert blocks == [
            ListItem(spans=[Span("first")], ordered=True),
            ListItem(spans=[Span("second")], ordered=True),
        ]

    def test_table(self):
        md = "| Name | Purpose |\n| --- | --- |\n| cookie_a | Session |\n| cookie_b | Analytics |"
        blocks = parse_markdown(md)
        assert blocks == [
            Table(
                headers=[[Span("Name")], [Span("Purpose")]],
                rows=[
                    [[Span("cookie_a")], [Span("Session")]],
                    [[Span("cookie_b")], [Span("Analytics")]],
                ],
            )
        ]

    def test_empty_content(self):
        assert parse_markdown("") == []
        assert parse_markdown("\n\n  \n") == []


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def make_policy(content: str, **kwargs) -> Policy:
    return Policy(
        tenant_id="ten_test456",
        tenant_name=kwargs.get("tenant_name", "Acme Inc"),
        title=kwargs.get("title", "Privacy Notice"),
        type=kwargs.get("type", PolicyType.PRIVACY_NOTICE),
        status=kwargs.get("status", PolicyStatus.DRAFT),
        content_format=PolicyContentFormat.MARKDOWN,
        current_version=kwargs.get("current_version", 1),
        content=content,
    )


SAMPLE_CONTENT = (
    "# Privacy Notice\n\n"
    "## Who We Are\n\n"
    "We are **Acme Inc**, located in Ireland. _Contact us anytime._\n\n"
    "## Your Rights\n\n"
    "- Right to access\n"
    "- Right to erasure\n\n"
    "## Cookies in Detail\n\n"
    "| Name | Purpose | Duration |\n"
    "| --- | --- | --- |\n"
    "| session_id | Session management | Session |\n"
    "| analytics | Usage analytics | 30 days |\n"
)


# ═══════════════════════════════════════════════════════════════════════════════
# generate_policy_pdf
# ═══════════════════════════════════════════════════════════════════════════════

class TestGeneratePolicyPdf:

    def test_returns_valid_pdf_bytes(self):
        pdf_bytes = generate_policy_pdf(make_policy(SAMPLE_CONTENT))
        assert pdf_bytes.startswith(b"%PDF")
        assert len(pdf_bytes) > 500

    def test_handles_empty_content(self):
        pdf_bytes = generate_policy_pdf(make_policy(""))
        assert pdf_bytes.startswith(b"%PDF")


# ═══════════════════════════════════════════════════════════════════════════════
# generate_policy_docx
# ═══════════════════════════════════════════════════════════════════════════════

class TestGeneratePolicyDocx:

    def test_returns_valid_docx(self):
        docx_bytes = generate_policy_docx(make_policy(SAMPLE_CONTENT))
        assert docx_bytes.startswith(b"PK")

        doc = DocxDocument(io.BytesIO(docx_bytes))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Who We Are", "Your Rights", "Cookies in Detail"]

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert [c.text for c in table.rows[0].cells] == ["Name", "Purpose", "Duration"]
        assert [c.text for c in table.rows[1].cells] == ["session_id", "Session management", "Session"]
        assert [c.text for c in table.rows[2].cells] == ["analytics", "Usage analytics", "30 days"]

    def test_handles_empty_content(self):
        docx_bytes = generate_policy_docx(make_policy(""))
        assert docx_bytes.startswith(b"PK")
